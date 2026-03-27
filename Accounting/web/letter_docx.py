"""
Letter DOCX Generator

Produces a formatted .docx letter using Response.docx as the master letterhead
template.  The header (company logo / full-page letterhead background) is
preserved from that template automatically.  The company stamp is embedded in
the body.  Fallback order:
  1. Response.docx  (preferred — full letterhead + logo)
  2. Templates.docx (legacy)
  3. Blank document with basic margins
"""

import base64
import io
import logging
import os
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

_REPO_ROOT     = Path(__file__).parent.parent
_RESPONSE_DOCX = _REPO_ROOT / "Response.docx"        # master letterhead template
_TEMPLATE_DOCX = _REPO_ROOT / "Templates.docx"        # legacy fallback
_ASSETS_DIR    = Path(__file__).parent / "data" / "letters" / "assets"
_OUTPUT_DIR    = Path(__file__).parent / "data" / "letters" / "docx"
_STAMP_PATH    = _ASSETS_DIR / "stamp.png"

_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
_ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def _ensure_stamp() -> bool:
    """Extract the company stamp from Response.docx into the assets folder."""
    if _STAMP_PATH.exists():
        return True
    if not _RESPONSE_DOCX.exists():
        return False
    try:
        with zipfile.ZipFile(str(_RESPONSE_DOCX)) as z:
            # image1.png is the stamp in the body of Response.docx
            data = z.read("word/media/image1.png")
            _STAMP_PATH.write_bytes(data)
            logger.info("Extracted stamp image to %s", _STAMP_PATH)
        return True
    except Exception as e:
        logger.warning("Could not extract stamp from Response.docx: %s", e)
        return False


def _remove_table_borders(table):
    """Remove all visible borders from a table (used for invisible layout tables)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "none")
        elem.set(qn("w:sz"), "0")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "auto")
        tblBorders.append(elem)
    tblPr.append(tblBorders)


def _add_sig_image(para, data_url: str, width_cm: float = 3.5):
    """Insert a base-64 signature image into a paragraph."""
    try:
        from docx.shared import Cm
        _header, b64data = data_url.split(",", 1)
        img_bytes = base64.b64decode(b64data)
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
    except Exception as e:
        logger.warning("Could not embed signature image: %s", e)
        para.add_run("[Signature image unavailable]")


def generate_letter_docx(letter: dict, signatures: dict) -> Optional[str]:
    """
    Generate a .docx file for the given letter, using the Response.docx
    letterhead template.

    Args:
        letter:     Letter record dict from letter_data_store.
        signatures: {role: {data_url, saved_at, ...}} from get_all_signatures().

    Returns:
        Absolute path to the generated .docx file, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        RIGHT  = WD_ALIGN_PARAGRAPH.RIGHT
        CENTER = WD_ALIGN_PARAGRAPH.CENTER
        LEFT   = WD_ALIGN_PARAGRAPH.LEFT

        # ── Load template ────────────────────────────────────────────────
        # Response.docx is the preferred template: its header contains the
        # full-page letterhead/logo background.  We clear only the body so
        # the header is preserved exactly as-is.
        if _RESPONSE_DOCX.exists():
            doc = Document(str(_RESPONSE_DOCX))
            # Remove all body content (paragraphs, tables) — sectPr is kept
            # so the header/footer link survives.
            body = doc.element.body
            for child in list(body):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("p", "tbl", "sdt"):
                    body.remove(child)
        elif _TEMPLATE_DOCX.exists():
            doc = Document(str(_TEMPLATE_DOCX))
            for para in doc.paragraphs:
                para.clear()
            while len(doc.paragraphs) > 1:
                doc.paragraphs[-1]._element.getparent().remove(
                    doc.paragraphs[-1]._element
                )
        else:
            doc = Document()
            for section in doc.sections:
                section.top_margin    = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin   = Cm(2.54)
                section.right_margin  = Cm(2.54)

        # ── Paragraph helper ─────────────────────────────────────────────
        def _para(text="", bold=False, italic=False, underline=False,
                  size=11, align=LEFT, space_after=6, color=None):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            if text:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            return p

        # ── Spacers (leave room below the letterhead header image) ────────
        # Response.docx has 4 empty paragraphs before the date line.
        for _ in range(4):
            _para(space_after=2)

        # ── Date & Reference (right-aligned, matching Response.docx) ─────
        letter_date = letter.get("date") or datetime.now().strftime("%d/%m/%Y")
        _para(f"Date: {letter_date}", size=11, align=RIGHT, space_after=2)
        _para(f"Ref. No: {letter.get('ref_number', '')}", size=11, align=RIGHT, space_after=6)
        _para(align=RIGHT, space_after=4)

        # ── Recipient ─────────────────────────────────────────────────────
        if letter.get("to"):
            _para(letter["to"], bold=True, size=11, space_after=2)
        if letter.get("to_address"):
            for line in letter["to_address"].splitlines():
                _para(line, size=10, space_after=2)
        _para(space_after=6)

        # ── Subject (centred, bold, underlined — matching Response.docx) ──
        _para(
            letter.get("subject", ""),
            bold=True, underline=True, size=11,
            align=CENTER, space_after=10,
        )

        # ── Salutation ────────────────────────────────────────────────────
        _para("Dear Sir / Madam,", size=11, space_after=8)

        # ── Body ──────────────────────────────────────────────────────────
        for block in (letter.get("body") or "").split("\n\n"):
            stripped = block.strip()
            if stripped:
                _para(stripped, size=11, space_after=8)

        # ── Closing ───────────────────────────────────────────────────────
        _para("Thank you for your time and consideration.", size=11, space_after=6)
        _para(space_after=4)

        # ── Stamp + "Best Regards" row (invisible 2-column table) ─────────
        # Left column: company stamp image
        # Right column: "Best Regards" closing text
        _ensure_stamp()
        stamp_table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(stamp_table)

        # Set column widths (5cm stamp, 9cm closing)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tblGrid = OxmlElement("w:tblGrid")
        for w in (2835, 5102):   # twips: 2835≈5cm, 5102≈9cm
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tblGrid.append(gc)
        stamp_table._tbl.insert(1, tblGrid)

        # Left cell — stamp
        left_cell = stamp_table.cell(0, 0)
        left_cell.paragraphs[0].clear()
        if _STAMP_PATH.exists():
            run = left_cell.paragraphs[0].add_run()
            run.add_picture(str(_STAMP_PATH), width=Cm(4.3), height=Cm(4.5))
        else:
            left_cell.paragraphs[0].add_run("[Company Stamp]").font.size = Pt(10)

        # Right cell — closing
        right_cell = stamp_table.cell(0, 1)
        right_cell.paragraphs[0].clear()
        right_cell.paragraphs[0].alignment = RIGHT
        br = right_cell.paragraphs[0].add_run("Best Regards,")
        br.bold = True
        br.font.size = Pt(11)

        # ── Authorised Signatures (PM / FM / MD) ──────────────────────────
        _para(space_after=8)
        _para("Authorised Signatures:", bold=True, size=10, space_after=4)

        auth_table = doc.add_table(rows=3, cols=3)
        auth_table.style = "Table Grid"
        col_labels = ["Project Manager (PM)", "Finance Manager (FM)", "Managing Director (MD)"]
        role_keys  = ["PM", "FM", "MD"]

        # Row 0: role titles
        for j, label in enumerate(col_labels):
            cell = auth_table.cell(0, j)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.size = Pt(9)

        # Row 1: signature images (or blank lines)
        for j, role in enumerate(role_keys):
            cell = auth_table.cell(1, j)
            cell.paragraphs[0].clear()
            sig = signatures.get(role)
            if sig and sig.get("data_url"):
                _add_sig_image(cell.paragraphs[0], sig["data_url"], width_cm=3.5)
            else:
                cell.paragraphs[0].add_run("_________________________").font.size = Pt(9)

        # Row 2: signed-by info
        for j, role in enumerate(role_keys):
            cell = auth_table.cell(2, j)
            cell.paragraphs[0].clear()
            letter_sigs = letter.get("signatures", {})
            if role in letter_sigs:
                info = letter_sigs[role]
                ts   = info.get("signed_at", "")
                try:
                    ts = datetime.fromisoformat(ts).strftime("%d %b %Y %H:%M")
                except Exception:
                    pass
                txt = f"Signed: {info.get('signed_by', '')}\n{ts}"
            else:
                txt = "Not yet signed"
            cell.paragraphs[0].add_run(txt).font.size = Pt(8)

        # ── CC ────────────────────────────────────────────────────────────
        if letter.get("cc"):
            _para(space_after=8)
            _para(f"CC: {letter['cc']}", size=10, color=(100, 100, 100), space_after=4)

        # ── Save ──────────────────────────────────────────────────────────
        ref = letter.get("ref_number") or letter["letter_id"]
        out_path = _OUTPUT_DIR / f"letter_{ref}.docx"
        doc.save(str(out_path))
        logger.info("Generated letter docx: %s", out_path)
        return str(out_path)

    except Exception as e:
        logger.error("generate_letter_docx failed: %s", e, exc_info=True)
        return None
